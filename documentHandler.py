import base64
import hashlib
import io
from datetime import datetime

import cv2
import numpy as np
from docx import Document
from docx.shared import Cm


def create_document(quality_criteria, img_string):
    document = Document()

    now = datetime.now()
    dt_string = now.strftime("%d/%m/%Y %H:%M:%S")

    document.add_heading("Отчет об оценке покрытия от " + dt_string, 0)

    r_data = base64.b64decode(img_string)
    img_as_np = np.frombuffer(r_data, dtype=np.uint8)
    img = cv2.imdecode(img_as_np, flags=1)

    cv2.imwrite("./temp.png", img)
    document.add_picture("./temp.png", width=Cm(10))

    document.add_heading("Критерии качества", level=2)

    table = document.add_table(rows=1, cols=2)

    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = "Критерий"
    hdr_cells[1].text = "Значение"

    row_cells = table.add_row().cells
    row_cells[0].text = "Максимальная толщина покрытия, мкм"
    row_cells[1].text = str(quality_criteria["max-coat-thickness"])

    row_cells = table.add_row().cells
    row_cells[0].text = "Медиана толщины покрытия, мкм"
    row_cells[1].text = str(quality_criteria["median-coat-thickness"])

    row_cells = table.add_row().cells
    row_cells[0].text = "Минимальная толщины покрытия, мкм"
    row_cells[1].text = str(quality_criteria["min-coat-thickness"])

    row_cells = table.add_row().cells
    row_cells[0].text = "Корреляция положения пор и глубины"
    row_cells[1].text = str(quality_criteria["pore-depth-correlation"])

    row_cells = table.add_row().cells
    row_cells[0].text = "Средняя площадь пор, мкм"
    row_cells[1].text = str(quality_criteria["avg-pore-square"])

    row_cells = table.add_row().cells
    row_cells[0].text = "Корреляция нерасплавленных частиц и глубины"
    row_cells[1].text = str(quality_criteria["unmelted-depth-correlation"])

    row_cells = table.add_row().cells
    row_cells[0].text = "Средняя площадь нерасплавленных частиц, мкм"
    row_cells[1].text = str(quality_criteria["avg-unmelt-square"])

    row_cells = table.add_row().cells
    row_cells[0].text = "Пористость, %"
    row_cells[1].text = str(quality_criteria["pores-percentage"])

    row_cells = table.add_row().cells
    row_cells[0].text = "Непроплавы, %"
    row_cells[1].text = str(quality_criteria["unmelts-percentage"])

    for cell in table.columns[0].cells:
        cell.width = Cm(10)

    # document.add_page_break()
    file = io.BytesIO()
    file.seek(0)

    sha1 = hashlib.sha1()
    sha = hashlib.sha1(("report" + dt_string).encode("utf-8")).hexdigest()
    # sha = hashlib.sha1('report' + dt_string)
    # sha.hexdigest()
    filename = sha + ".docx"
    document.save("./reports/" + filename)
    return filename
